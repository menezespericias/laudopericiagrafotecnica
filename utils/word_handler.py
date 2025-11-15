from docx import Document
from docx.shared import Inches
from num2words import num2words
import os
import re

# 1. FUNÇÃO PRINCIPAL DE GERAÇÃO
def gerar_laudo(caminho_modelo, caminho_saida, dados, anexos, adendos):
    doc = Document(caminho_modelo)
    
    # Adicionar versões por extenso dos números para uso automático
    if "NUM_ESPECIMES" in dados and dados["NUM_ESPECIMES"].isdigit():
        num = int(dados["NUM_ESPECIMES"])
        dados["NUM_ESPECIMES_EXTENSO"] = num2words(num, lang='pt_BR')

    # Substituição de campos de texto em parágrafos e tabelas (PRESERVANDO FORMATAÇÃO)
    for paragrafo in doc.paragraphs:
        substituir_em_paragrafo(paragrafo, dados)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                # Iterar sobre os parágrafos dentro de cada célula
                for paragrafo in celula.paragraphs:
                    substituir_em_paragrafo(paragrafo, dados)

    # Inserção de listas (AUTORES, REUS)
    # ATENÇÃO: Se os placeholders forem removidos, esta função não fará nada.
    # É preciso garantir que ela seja chamada antes do tratamento geral de placeholders se o seu modelo
    # tiver os placeholders de lista como parágrafos inteiros.
    # No código main.py, a lista é tratada como um dado simples. Vamos garantir que se for necessário
    # formatar como lista, o faça aqui.
    
    inserir_lista_no_paragrafo(doc, "AUTORES", dados["AUTORES"])
    inserir_lista_no_paragrafo(doc, "REUS", dados["REUS"])


    # Inserção de imagens das assinaturas (adicionado cabeçalho)
    doc.add_page_break()
    doc.add_heading("IX. ANÁLISE DAS ASSINATURAS", level=1)
    
    for assinatura in dados["ASSINATURAS"]:
        if assinatura["nome"]:
            doc.add_paragraph(f"Assinatura Analisada: {assinatura['nome']}", style='Heading 2')
        if assinatura["observacoes"]:
            doc.add_paragraph(f"Observações: {assinatura['observacoes']}")
        
        # O objeto do Streamlit é um UploadedFile. 
        # doc.add_picture aceita o objeto de arquivo aberto diretamente.
        if assinatura["tabela"] is not None:
             # Reposiciona o ponteiro do arquivo para o início antes de ler
            assinatura["tabela"].seek(0)
            doc.add_picture(assinatura["tabela"], width=Inches(5.5))
            doc.add_paragraph() # Espaçamento

    # Inserção de anexos
    if anexos:
        doc.add_page_break()
        doc.add_heading("X. ANEXOS", level=1)
        for imagem in anexos:
            if imagem is not None:
                imagem.seek(0)
                doc.add_picture(imagem, width=Inches(5.5))
                doc.add_paragraph(f"Figura: {imagem.name}") # Adiciona legenda simples
                doc.add_paragraph()

    # Inserção de adendos
    if adendos:
        doc.add_page_break()
        doc.add_heading("XI. ADENDOS", level=1)
        for imagem in adendos:
            if imagem is not None:
                imagem.seek(0)
                doc.add_picture(imagem, width=Inches(5.5))
                doc.add_paragraph(f"Figura: {imagem.name}") # Adiciona legenda simples
                doc.add_paragraph()

    # Cria diretório de saída se não existir
    os.makedirs(os.path.dirname(caminho_saida), exist_ok=True)
    doc.save(caminho_saida)

# 2. FUNÇÃO CRÍTICA DE SUBSTITUIÇÃO (PRESERVA FORMATAÇÃO)
def substituir_em_paragrafo(paragrafo, dados):
    """Substitui placeholders em um parágrafo, preservando a formatação (runs)."""
    
    # 1. Agrega o texto de todos os runs do parágrafo
    texto_completo = "".join(run.text for run in paragrafo.runs)
    
    # 2. Executa a substituição em todo o texto
    texto_substituido = substituir_placeholders(texto_completo, dados)
    
    # 3. Se o texto não mudou, retorna
    if texto_completo == texto_substituido:
        return
    
    # 4. Limpa o parágrafo existente (deleta runs antigos)
    # Manteremos o primeiro run e apagaremos os outros
    if not paragrafo.runs:
        paragrafo.add_run(texto_substituido)
        return

    primeiro_run = paragrafo.runs[0]
    primeiro_run.text = texto_substituido
    
    # Remove todos os runs a partir do segundo
    for i in range(len(paragrafo.runs) - 1, 0, -1):
        p = paragrafo._element
        p.remove(paragrafo.runs[i]._element)

# 3. FUNÇÃO AUXILIAR DE SUBSTITUIÇÃO DE TEXTO
def substituir_placeholders(texto, dados):
    """Substitui placeholders (ex: [CHAVE]) pelo valor em dados."""
    
    # Padrões para placeholders: [CHAVE], {CHAVE}, <<CHAVE>>
    # O r-string é uma "raw string" para facilitar o regex
    padrao = r"(\[\s*([^\]]+?)\s*\]|\{\s*([^}]+?)\s*\}|<<\s*([^>]+?)\s*>>)"
    
    def replace_match(match):
        # O grupo que não for None contém a CHAVE
        chave = match.group(2) or match.group(3) or match.group(4)
        chave_limpa = chave.strip()

        # Tratamento para versão por extenso
        if chave_limpa.endswith("_EXTENSO"):
            base_chave = chave_limpa[:-8] # Remove '_EXTENSO'
            if base_chave in dados and isinstance(dados[base_chave], str) and dados[base_chave].isdigit():
                num = int(dados[base_chave])
                return num2words(num, lang='pt_BR')
        
        # Tratamento padrão
        if chave_limpa in dados:
            valor = dados[chave_limpa]
            if isinstance(valor, str):
                return valor
            # Se for uma lista, retorna a união dos itens (caso o placeholder tenha sido esquecido pela inserção de lista)
            elif isinstance(valor, list):
                 return "\n".join(valor)
        
        # Se não encontrou valor ou o valor não é string, retorna o placeholder original
        return match.group(0)

    # Realiza a substituição usando a função interna (callback)
    return re.sub(padrao, replace_match, texto)


# 4. FUNÇÃO PARA INSERIR LISTA (Como Parágrafos)
def inserir_lista_no_paragrafo(doc, marcador, lista):
    """
    Localiza o placeholder do marcador no documento e o substitui por itens da lista
    como novos parágrafos formatados com uma lista simples.
    """
    placeholder = f"[{marcador}]"
    
    for i, paragrafo in enumerate(doc.paragraphs):
        if placeholder in paragrafo.text:
            # Encontrou o placeholder, vamos criar os novos parágrafos AQUI.
            
            # Remove o placeholder do parágrafo original
            paragrafo.text = paragrafo.text.replace(placeholder, "")
            
            if not lista:
                return # Se a lista estiver vazia, apenas remove o placeholder

            # Adiciona os novos itens da lista como novos parágrafos logo após o original
            parent_element = paragrafo._element.getparent()
            
            for item in lista:
                # Cria um novo parágrafo com o estilo de lista de bullet points
                novo_paragrafo = parent_element.addprevious(paragrafo._element.tag, item)
                novo_paragrafo.text = item
                
                # Aplica estilo de lista (você pode precisar ajustar o estilo dependendo do seu modelo)
                try:
                    novo_paragrafo.style = 'List Bullet' 
                except KeyError:
                    # Tenta um estilo padrão mais comum se 'List Bullet' não estiver no modelo
                    try:
                        novo_paragrafo.style = 'List Paragraph'
                    except KeyError:
                        pass # Continua sem estilo específico

            # O placeholder foi removido e a lista inserida. 
            # Como a função 'substituir_em_paragrafo' é mais robusta,
            # vamos manter a lógica de lista simples na principal, mas esta seria a forma de fazer 
            # uma lista real com bullet points.
            break

# A função original 'inserir_lista' está obsoleta, use 'substituir_placeholders' 
# junto com 'substituir_em_paragrafo' que agora pode tratar listas internamente.
# Para manter a lista de Autores/Réus em parágrafos separados, remova esta linha abaixo
# e use o 'inserir_lista_no_paragrafo' adaptado.

# A maneira mais simples (e a que usei no código final) é deixar a substituição de lista
# ser tratada pelo 'substituir_placeholders' que usa '\n'.

def inserir_lista(doc, marcador, lista):
    # Esta função não é mais necessária com a nova implementação de 'substituir_em_paragrafo'
    # mas mantida aqui por compatibilidade.
    pass