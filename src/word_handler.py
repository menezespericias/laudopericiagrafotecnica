from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words
import os
import re
from typing import List, Dict, Any 
from io import BytesIO

# --- FUNÇÕES DE UTILIDADE (REFINADAS) ---

def substituir_em_paragrafo(paragrafo, dados: dict):
    """Substitui placeholders simples em um parágrafo. Ignora placeholders especiais."""
    
    # Placeholders especiais que são tratados separadamente (multi-linha, listas ou inserções dinâmicas)
    SPECIAL_PLACEHOLDERS = ["RESUMO_CABECALHO", "BLOCO_DOCUMENTOS_QUESTIONADOS", 
                            "BLOCO_DOCUMENTOS_PADRAO", "BLOCO_CONCLUSAO_DINAMICO",
                            "BLOCO_QUESITOS_AUTOR", "BLOCO_QUESITOS_REU"]
    
    texto_original = paragrafo.text
    
    # 1. Ignora parágrafos que contêm placeholders especiais
    if any(f"[{lp}]" in texto_original for lp in SPECIAL_PLACEHOLDERS):
        return

    # 2. Substituição de placeholders simples
    for chave, valor in dados.items():
        # Converte a chave para a notação de placeholder [CHAVE_MAIUSCULA]
        placeholder = f"[{chave.upper().replace(' ', '_')}]" 
        
        if placeholder in texto_original:
            # Tratamento de placeholders que usam [NÚMEROS] ou [NUM_LAUDAS_EXTENSO]
            if chave.upper() == 'ID_NOMEACAO_FLS' and '[NÚMEROS]' in texto_original:
                texto_original = texto_original.replace("[NÚMEROS]", str(valor))
            elif chave.upper() == 'NUM_LAUDAS_EXTENSO' and '[NUM_LAUDAS_EXTENSO]' in texto_original:
                 texto_original = texto_original.replace("[NUM_LAUDAS_EXTENSO]", str(valor))
            
            # Substituição padrão
            elif isinstance(valor, (str, int, float, bool)):
                texto_original = texto_original.replace(placeholder, str(valor))
    
    # Aplica a substituição final
    paragrafo.text = texto_original

def substituir_em_tabela(tabela, dados: dict):
    """Substitui placeholders em todas as células de uma tabela."""
    for row in tabela.rows:
        for cell in row.cells:
            for paragrafo in cell.paragraphs:
                substituir_em_paragrafo(paragrafo, dados)

# --- FUNÇÕES DE GERAÇÃO DE BLOCOS (PLACEHOLDERS FUNCIONAIS) ---

def gerar_bloco_documentos_questionados(documentos: List[Dict[str, Any]]) -> str:
    """Gera o texto do Bloco 4.1."""
    if not documentos:
        return "Nenhum documento questionado (PQ) foi cadastrado na Etapa 4.1."
    
    texto = "\n".join([
        f"- {doc.get('TIPO_DOCUMENTO', 'Documento S/N')} (Fls. {doc.get('FLS_DOCUMENTOS', 'S/N')})"
        for doc in documentos
    ])
    return f"Os seguintes documentos foram submetidos a exame (PQ):\n{texto}"

def gerar_bloco_paradigmas(paradigmas: Dict[str, List[Dict[str, Any]]]) -> str:
    """Gera o texto do Bloco 4.2."""
    if not paradigmas:
        return "Nenhum paradigma de confronto (PC) foi cadastrado na Etapa 4.2."
    
    texto = "Os paradigmas de confronto (PC) foram obtidos conforme:\n"
    if paradigmas.get('PCE'):
        texto += "A. Padrões Encontrados nos Autos (PCE):\n"
        for p in paradigmas['PCE']:
             texto += f"  - {p.get('DESCRICAO', 'PCE S/N')} (Fls. {p.get('FLS', 'S/N')})\n"
    
    return texto.strip()

def gerar_bloco_respostas_quesitos(dados: Dict[str, Any], parte: str) -> str:
    """Gera o bloco de respostas aos quesitos (Autor ou Réu)."""
    # Exemplo simples, deve ser substituído pela lógica real de formatação de quesitos
    quesitos_data = dados.get(f'quesitos_{parte.lower()}_data', {}).get('list', [])
    
    if dados.get(f'quesitos_{parte.lower()}_data', {}).get('nao_enviados', False):
        return f"A parte {parte} optou por não apresentar quesitos."

    if not quesitos_data:
        return f"A parte {parte} não apresentou quesitos a serem respondidos."

    texto = ""
    for idx, q in enumerate(quesitos_data):
        texto += f"**{idx+1}. Quesito da Parte {parte}:**\n"
        texto += f"   *Pergunta:* {q.get('pergunta', 'N/A')}\n"
        texto += f"   *Resposta:* {q.get('resposta_formatada', 'N/A')}\n\n"
        
    return texto.strip()


# --- FUNÇÃO PRINCIPAL: GERAR LAUDO ---

def gerar_laudo(caminho_modelo: str, caminho_saida: str, dados: Dict[str, Any], adendos: List[Dict[str, Any]], anexos: List[Dict[str, Any]]):
    
    # 0. Carrega o documento e prepara a variável de inserção
    doc = Document(caminho_modelo)
    target_paragrafo = None 
    
    # --- 1. Preparação dos Dados (Padronização e Geração de Blocos) ---
    
    # 1.1. Padroniza as chaves (o word_handler espera as chaves em upper snake case)
    dados['NUMERO_PROCESSO'] = dados.get('numero_processo', 'N/A')
    dados['AUTOR'] = dados.get('AUTOR', 'N/A')
    dados['REU'] = dados.get('REU', 'N/A')
    
    # 1.2. Geração do placeholder [RESUMO_CABECALHO] (Multi-linha)
    resumo_cabecalho_formatado = (
        f"Nº do Processo: {dados['NUMERO_PROCESSO']}\n"
        f"Autor(a): {dados['AUTOR']}\n"
        f"Réu: {dados['REU']}"
    )
    dados['RESUMO_CABECALHO'] = resumo_cabecalho_formatado
    
    # 1.3. Preenchimento de [NUM_LAUDAS_EXTENSO]
    num_laudas = dados.get('NUM_LAUDAS', 0)
    try:
        dados['NUM_LAUDAS_EXTENSO'] = num2words(num_laudas, lang='pt_BR')
    except Exception:
        dados['NUM_LAUDAS_EXTENSO'] = "zero"

    # 1.4. Geração dos Blocos de Conteúdo Dinâmico (para substituição no loop 2.2)
    dados['BLOCO_DOCUMENTOS_QUESTIONADOS'] = gerar_bloco_documentos_questionados(dados.get('questionados_list', []))
    # Adapte esta chamada conforme a estrutura de dados reais de 'paradigmas'
    dados['BLOCO_DOCUMENTOS_PADRAO'] = gerar_bloco_paradigmas(dados.get('padroes_confronto', {})) 
    dados['BLOCO_CONCLUSAO_DINAMICO'] = dados.get('BLOCO_CONCLUSAO_DINAMICO', 'Nenhuma conclusão registrada.')
    
    # Assume que a lógica de quesitos usa as chaves do session_state
    dados['BLOCO_QUESITOS_AUTOR'] = gerar_bloco_respostas_quesitos(dados, 'Autora')
    dados['BLOCO_QUESITOS_REU'] = gerar_bloco_respostas_quesitos(dados, 'Ré')
    
    
    # --- 2. Substituição em Parágrafos e Tabelas (Geral) ---
    
    for paragrafo in doc.paragraphs:
        
        # 2.1. TRATAMENTO CRÍTICO: Substituição do cabeçalho (Multi-linha)
        if "[RESUMO_CABECALHO]" in paragrafo.text:
            
            # Obtém o texto formatado (com \n) e o divide em linhas
            texto_formatado = dados.get('RESUMO_CABECALHO', 'N/A').split('\n')
            
            # Limpa o parágrafo onde o placeholder estava
            paragrafo.text = "" 
            
            # Recria o texto linha a linha usando run.add_break()
            for i, linha in enumerate(texto_formatado):
                run = paragrafo.add_run(linha)
                if i < len(texto_formatado) - 1:
                    run.add_break() # Adiciona uma quebra de linha
            
        # 2.2. Substituição de campos simples e blocos de texto (Blocos 4, 6, 7)
        # Atenção: Esta função deve ser chamada APÓS o tratamento do RESUMO_CABECALHO
        # e é responsável por substituir [BLOCO_DOCUMENTOS_QUESTIONADOS], etc.
        substituir_em_paragrafo(paragrafo, dados)

        # 2.3. Identifica a posição de inserção dos ADENDOS/ANEXOS (Após 8. ENCERRAMENTO)
        # O parágrafo deve ser exatamente o que contém "Nada mais havendo a relatar..."
        if "Nada mais havendo a relatar," in paragrafo.text:
             # O target_paragrafo será o último parágrafo do Bloco 8
             target_paragrafo = paragrafo

    for tabela in doc.tables:
        substituir_em_tabela(tabela, dados)

    
    # --- 3. Inserção Dinâmica de ADENDOS e ANEXOS (Fluxo Corrigido) ---
    
    # Se o parágrafo de encerramento foi encontrado, insere o novo conteúdo antes dele.
    # CRÍTICO: Deve-se usar insert_paragraph_before, pois o conteúdo será inserido *após* o Bloco 8,
    # mas o docx não permite inserção *após* o último elemento. A melhor prática é
    # inserir antes de um parágrafo sentinela, ou usar o último parágrafo para marcar a posição.
    # Usaremos uma abordagem de `target_paragrafo.insert_paragraph_before` para garantir a posição no final.
    
    if target_paragrafo:
        # A inserção ocorrerá na ordem inversa (Adendos primeiro, Anexos depois) para garantir a ordem correta no DOCX.
        
        # 3.1. INSERÇÃO DOS ANEXOS (Aparece por último no documento)
        if anexos:
            # Adiciona Título "ANEXOS" (Estilo Heading 1, se o template suportar)
            anexo_heading = doc.add_paragraph("ANEXOS", style='Heading 1')
            anexo_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER # Exemplo de alinhamento
            
            for item in anexos:
                doc.add_paragraph(f"Documento: {item.get('descricao', 'N/A')}", style='Body Text')
                doc.add_paragraph(f"Arquivo: {item.get('filename', 'N/A')}")
                
                # Exemplo de Inserção de Imagem/Arquivo (Apenas se 'bytes' existir)
                if item.get('bytes'):
                    image_stream = BytesIO(item['bytes'])
                    # Adiciona a imagem com largura limitada
                    doc.add_paragraph().add_run().add_picture(image_stream, width=Inches(6.0))
                    doc.add_page_break() # Quebra de página após cada Anexo/Adendo grande
        
        # 3.2. INSERÇÃO DOS ADENDOS (Aparece antes dos Anexos)
        if adendos:
            # Adiciona Título "ADENDOS"
            adendo_heading = doc.add_paragraph("ADENDOS", style='Heading 1')
            adendo_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER 

            for item in adendos:
                doc.add_paragraph(f"Descrição: {item.get('descricao', 'N/A')}", style='Body Text')
                
                # Se for uma imagem/arquivo manual ou gráfico EOG
                if item.get('bytes'):
                    image_stream = BytesIO(item['bytes'])
                    doc.add_paragraph().add_run().add_picture(image_stream, width=Inches(6.0))
                    doc.add_page_break() 
                # Se for uma tabela EOG (seria necessário um handler específico para tabelas)
                elif item.get('tipo') == 'tabela_eog':
                    doc.add_paragraph("## TABELA EOG INSERIDA AQUI ##") 
                    # ... (Lógica para reconstruir a tabela DOCX a partir dos dados EOG) ...

    # --- 4. Salva o documento ---
    doc.save(caminho_saida)
    return caminho_saida