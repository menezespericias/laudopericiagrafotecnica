from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words
import os
import re
from typing import List, Dict, Any 
from io import BytesIO

# --- FUNÇÕES DE UTILIDADE (REFINADAS) ---

def substituir_em_paragrafo(paragrafo, dados: dict):
    # Placeholders que serão tratados como LISTAS (não substituição simples)
    LIST_PLACEHOLDERS = ["AUTORES", "REUS"] 
    
    texto_original = paragrafo.text
    
    # Ignora parágrafos que contêm placeholders de lista (tratados em outra função)
    if any(f"[{lp}]" in texto_original for lp in LIST_PLACEHOLDERS):
        return

    for chave, valor in dados.items():
        # Trata chaves com espaço, convertendo-as para a notação de placeholder do template
        placeholder = f"[{chave.upper().replace(' ', '_')}]" 
        
        if placeholder in texto_original:
            # Substitui placeholders de FLS. com o mesmo texto (Ex: [NUMEROS] em 1. Apresentação)
            if chave.upper() == 'ID_NOMEACAO_FLS':
                texto_original = texto_original.replace("[NUMEROS]", str(valor))
                paragrafo.text = texto_original # Atualiza o parágrafo antes de quebrar para run
                
            # Verifica se o valor não é uma lista/dicionário (apenas substituição de texto simples)
            valor_str = str(valor) if not isinstance(valor, (list, dict)) else ""
            
            if valor_str:
                for run in paragrafo.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, valor_str)
                        
def substituir_em_tabela(tabela, dados: dict):
    for row in tabela.rows:
        for cell in row.cells:
            for paragrafo in cell.paragraphs:
                substituir_em_paragrafo(paragrafo, dados)

def inserir_imagens_em_secao(doc: Document, titulo: str, adendos: List[Dict[str, Any]], prefixo: str):
    """Insere o título da seção e as imagens dos adendos."""
    
    # Filtra adendos que possuem imagem_obj (BytesIO) e possuem dados
    imagens_para_inserir = [a for a in adendos if a.get("imagem_obj") and isinstance(a.get("imagem_obj"), BytesIO) and a.get("imagem_obj").getbuffer().nbytes > 0]
    
    if not imagens_para_inserir:
        return
        
    doc.add_heading(titulo, level=2)
    
    for i, adendo in enumerate(imagens_para_inserir):
        legenda = adendo.get("legenda", f"{prefixo} {i+1}")
        file_obj_buffer = adendo.get("imagem_obj") # É um BytesIO
        
        try:
            # Garante que o buffer esteja no início antes de ler
            file_obj_buffer.seek(0) 
            
            # Insere a imagem usando o BytesIO
            doc.add_picture(file_obj_buffer, width=Inches(6))
            
            # Adiciona a legenda
            p = doc.add_paragraph(legenda)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            # Em caso de erro com a imagem 
            doc.add_paragraph(f"ERRO: Não foi possível inserir a imagem: {legenda}. {e}")

# --- FUNÇÕES DE GERAÇÃO DE CONTEÚDO DINÂMICO ---

def gerar_bloco_analise_dinamica(doc: Document, dados: dict):
    """
    Gera o bloco de texto dinâmico para o 5.2. Confronto Grafoscópico (Analise dos EOGs).
    Substitui a lista padronizada do template pela lista filtrada.
    """
    # 1. Prepara a lista de EOGs analisados
    eog_data = dados.get('EOG_ANALYSIS', {})
    
    # Filtra apenas Convergente e Divergente
    analise_list = [
        f"**{element}**: {conclusion}"
        for element, conclusion in eog_data.items() 
        if conclusion in ["Convergente", "Divergente"]
    ]
    
    # 2. Encontra a seção 5.2 no template para inserir a lista de EOGs
    paragrafo_ref_52 = None
    
    # Usamos uma lista de referências do template para encontrar o ponto de inserção/substituição
    ref_list = [
        "Natureza do Gesto Gráfico: Velocidade, pressão, espontaneidade.", 
        "Valores Angulares e Curvilíneos: Inclinação dos traços."
    ]
    
    for p in doc.paragraphs:
        if any(ref in p.text for ref in ref_list):
            paragrafo_ref_52 = p
            break
            
    if not paragrafo_ref_52:
        # Tenta encontrar o placeholder do ANALISE_TEXTO se a lista padrão falhar (fallback)
        for p in doc.paragraphs:
            if "[ANALISE_TEXTO]" in p.text:
                p.text = p.text.replace("[ANALISE_TEXTO]", dados.get('ANALISE_TEXTO', ''))
                return 

    # 3. Remove os itens padrões do template e insere os novos itens dinâmicos
    if paragrafo_ref_52:
        paragrafos_a_remover = []
        # Coleta os parágrafos que contêm os itens padrões
        for p in doc.paragraphs:
            for ref in ref_list:
                if ref in p.text:
                    paragrafos_a_remover.append(p)
                    
        # Remove o texto dos itens padrões
        for p_rem in paragrafos_a_remover:
            p_rem.text = "" 

        # Insere os novos itens dinâmicos (Convergente/Divergente)
        for item in reversed(analise_list):
            novo_paragrafo = paragrafo_ref_52.insert_paragraph_before(f"- {item}")
            
        # Adiciona o texto livre (ANALISE_TEXTO) logo após o 5.2
        paragrafo_ref_52.insert_paragraph_before("")
        paragrafo_ref_52.insert_paragraph_before(dados.get('ANALISE_TEXTO', ''))


def gerar_bloco_conclusao_dinamico(dados: dict) -> str:
    """Gera o bloco de conclusão formatado em negrito com o tipo principal."""
    tipo = dados.get('CONCLUSÃO_TIPO', 'Inconclusiva')
    texto = dados.get('CONCLUSION', '')
    
    if tipo == "Selecione a Conclusão":
        return ""

    # Usa o formato de negrito e caixa alta no tipo de conclusão
    return f"**{tipo.upper()}**.\n{texto}"

def gerar_bloco_respostas_quesitos(dados: dict, parte: str) -> str:
    """Gera o bloco de respostas aos quesitos de uma parte (Autor ou Réu)."""
    respostas = dados.get('RESPOSTAS_QUESITOS_LIST', [])
    
    bloco_texto = []
    
    # Filtra as respostas para a parte específica
    respostas_filtradas = [r for r in respostas if r['parte'] == parte]
    
    if not respostas_filtradas:
        return f"Não foram apresentados quesitos pela parte {parte}."

    for r in respostas_filtradas:
        # Formato: X.1. Quesito: [Texto do Quesito] Resposta: [Texto da Resposta]
        bloco_texto.append(f"**Quesito {r['id']}:** {r['quesito']}\n**Resposta:** {r['resposta']}\n")
        
    # Converte a lista de strings formatadas em uma única string com quebras de linha
    return "\n".join(bloco_texto)

def gerar_bloco_documentos_questionados(doc: Document, dados: dict):
    """Gera a lista de Documentos Questionados (4.1) no corpo do laudo, substituindo os placeholders 0 a 5."""
    docs = dados.get('DOCUMENTOS_QUESTIONADOS', [])
    
    for i in range(6):
        # A numeração no template geralmente começa em 1, mas os placeholders podem estar em [0]
        placeholder_tipo = f"[TIPO_DOCUMENTO_{i}]"
        placeholder_num = f"[NÚMERO DO CONTRATO_{i}]"
        placeholder_data = f"[DATA_DOCUMENTO_{i}]" if i == 0 else f"[DATA_{i}]"
        placeholder_fls = f"[FLS_DOCUMENTOS_{i}]" if i == 0 else f"[NÚMEROS DAS FOLHAS_{i}]"
        
        texto_doc = ""
        if i < len(docs):
            d = docs[i]
            # O texto completo que substitui o bloco da linha Documento X:
            texto_doc = f"{d['tipo']}, nº {d['numero']}, datado de {d['data']}, fls. {d['fls']}."
        
        # Procura e substitui todos os placeholders para este documento
        for paragrafo in doc.paragraphs:
            if placeholder_tipo in paragrafo.text or placeholder_num in paragrafo.text:
                if texto_doc:
                    # Tenta substituir o bloco completo (Documento X: [TIPO_DOCUMENTO_X]...)
                    padrao_completo = re.compile(r"Documento\s+" + str(i+1) + r":\s*" + re.escape(placeholder_tipo) + r".*")
                    
                    if padrao_completo.search(paragrafo.text):
                        # Se encontrar o padrão, substitui a linha inteira
                        paragrafo.text = re.sub(padrao_completo, f"Documento {i+1}: {texto_doc}", paragrafo.text)
                    else:
                        # Se não encontrar o padrão completo, faz a substituição individual (mais robusto)
                        paragrafo.text = paragrafo.text.replace(placeholder_tipo, d['tipo'])
                        paragrafo.text = paragrafo.text.replace(placeholder_num, d['numero'])
                        paragrafo.text = paragrafo.text.replace(placeholder_data, d['data'])
                        paragrafo.text = paragrafo.text.replace(placeholder_fls, d['fls'])
                else:
                    # Remove a linha se não houver documento
                    padrao_completo = re.compile(r"Documento\s+" + str(i+1) + r":\s*" + re.escape(placeholder_tipo) + r".*")
                    if padrao_completo.search(paragrafo.text):
                        paragrafo.text = ""
                    else:
                        paragrafo.text = paragrafo.text.replace(placeholder_tipo, "")
                        paragrafo.text = paragrafo.text.replace(placeholder_num, "")
                        paragrafo.text = paragrafo.text.replace(placeholder_data, "")
                        paragrafo.text = paragrafo.text.replace(placeholder_fls, "")

def gerar_bloco_padroes_encontrados(doc: Document, dados: dict):
    """Gera a lista de Padrões Encontrados nos Autos (4.2.B)."""
    placeholder_base = "[TIPO_DOCUMENTO] – Fls. [NÚMEROS], datado de [DATA]."
    padroes = dados.get('PADROES_ENCONTRADOS', [])
    
    # Encontra o parágrafo que contém o placeholder do Documento 1 em 4.2.B
    paragrafo_encontrado = None
    paragrafos_a_remover = []
    
    for p in doc.paragraphs:
        if "Documento 1:" in p.text and placeholder_base in p.text:
            paragrafo_encontrado = p
        elif "Documento 2:" in p.text and placeholder_base in p.text:
            paragrafos_a_remover.append(p)

    if not paragrafo_encontrado:
        return

    # Remove os parágrafos de Documento 2 e 3 (se existirem)
    for p_rem in paragrafos_a_remover:
        p_rem.text = ""

    for i, p in enumerate(padroes):
        texto_padrao = f"Documento {i+1}: {p['tipo']} – Fls. {p['fls']}, datado de {p['data']}."
        
        if i == 0:
            paragrafo_encontrado.text = texto_padrao
        else:
            paragrafo_encontrado.insert_paragraph_before(texto_padrao)


# --- FUNÇÃO PRINCIPAL DE GERAÇÃO ---

def gerar_laudo(
    caminho_modelo: str, 
    caminho_saida: str, 
    dados: dict, 
    adendos: List[Dict[str, Any]] 
):
    
    doc = Document(caminho_modelo)
        
    # --- 1. Pré-Cálculos e Extensos ---
    
    # Cálculos para campos por extenso
    for campo in ["NUM_ESPECIMES", "NUM_LAUDAS"]:
        valor = dados.get(campo, "0")
        if isinstance(valor, str) and valor.isdigit():
            num = int(valor)
            dados[f"{campo}_EXTENSO"] = num2words(num, lang='pt_BR').upper() 
        else:
            dados[f"{campo}_EXTENSO"] = ""

    # Campos de texto livre que viram lista
    dados['AUTORES'] = [a.strip() for a in dados.get('autor', '').split('\n') if a.strip()]
    dados['REUS'] = [r.strip() for r in dados.get('reu', '').split('\n') if r.strip()]
    
    # Define os primeiros nomes e substitui placeholders de nome no dicionário para substituição simples
    dados['PRIMEIRO_AUTOR'] = dados['AUTORES'][0] if dados['AUTORES'] else "Autor(a) Não Informado(a)"
    dados['NOME COMPLETO DO RÉU'] = dados['REUS'][0] if dados['REUS'] else "Réu Não Informado"
    dados['NOME DO AUTOR'] = dados['PRIMEIRO_AUTOR']
    dados['NOME COMPLETO DO PERITO'] = dados.get('PERITO', '')
    dados['NÚMERO DE REGISTRO'] = dados.get('NUMERO_REGISTRO', '')
    
    # Bloco de conclusão dinâmico (tratado como substituição de texto simples)
    dados['BLOCO_CONCLUSAO_DINAMICO'] = gerar_bloco_conclusao_dinamico(dados)
    
    # Bloco de respostas aos quesitos
    dados['BLOCO_QUESITOS_AUTOR'] = gerar_bloco_respostas_quesitos(dados, 'Autor')
    dados['BLOCO_QUESITOS_REU'] = gerar_bloco_respostas_quesitos(dados, 'Réu')
    
    # --- 2. Substituição em Parágrafos e Tabelas (Geral) ---
    
    for paragrafo in doc.paragraphs:
        
        # 2.1. Substituição do cabeçalho (Resumo)
        if "[NUMERO DO PROCESSO]Autor: [nome do autor]Réu: [nome do réu]" in paragrafo.text:
            texto_resumo = f"{dados.get('numero_processo', '')} Autor: {dados.get('PRIMEIRO_AUTOR', '')} Réu: {dados.get('NOME COMPLETO DO RÉU', '')}"
            paragrafo.text = paragrafo.text.replace(f"Processo: [NUMERO DO PROCESSO]Autor: [nome do autor]Réu: [nome do réu]", f"Processo: {texto_resumo}")
        
        # 2.2. Substituição de campos simples 
        substituir_em_paragrafo(paragrafo, dados)

    for tabela in doc.tables:
        substituir_em_tabela(tabela, dados)

    # --- 3. Inserção de Listas e Conteúdo Estruturado ---
    
    # 3.1. Inserção dos Documentos Questionados (4.1)
    gerar_bloco_documentos_questionados(doc, dados)
    
    # 3.2. Inserção dos Padrões Encontrados (4.2.B)
    gerar_bloco_padroes_encontrados(doc, dados)

    # 3.3. Geração e Inserção da Análise Detalhada de EOG (5.2)
    gerar_bloco_analise_dinamica(doc, dados)
    
    # 3.4. Inserção de Adendos Gráficos (Bloco 6) - Inclui Tabela e Gráfico EOG
    inserir_imagens_em_secao(doc, "6. ADENDOS GRÁFICOS", adendos, "Adendo")

    doc.save(caminho_saida)